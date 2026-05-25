import base64
import io
import logging
import re
import time
from typing import List, Optional

import cv2
import numpy as np
import pytesseract
from docx import Document
from docx.shared import Pt
from fastapi import FastAPI, File, HTTPException, Query, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response
from PIL import Image as PILImage
from pydantic import BaseModel, Field

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

OCR_LANG = "amh+eng"
ETHIOPIC_RE = re.compile(r"[\u1200-\u137F]")
LATIN_WORD_RE = re.compile(r"[a-zA-Z]{2,}")
ETHIOPIC_WORD_RE = re.compile(r"[\u1200-\u137F]{2,}")
REAL_CHAR_RE = re.compile(r"[a-zA-Z\u1200-\u137F\d]")
DOT_RUN_RE = re.compile(r"[.\-_·•]{3,}")
SPACE_RUN_RE = re.compile(r"  +")
MAX_BATCH_PAGES = 50
DOCX_FONT = "Nyala"

# Per-word confidence floor (0-100). Amharic LSTM scores are conservative; 25-30 is realistic.
MIN_WORD_CONF = 30
# Average per-line confidence required to keep a line. Filters trailing edge noise.
MIN_LINE_CONF = 55
# Long edge to upscale to before OCR. Higher = better recognition but slower.
OCR_TARGET_PX = 2400
# Long edge for pipeline preview images (fast to encode).
PREVIEW_PX = 1200
# Minimum % of the image area a detected quadrilateral must cover to be trusted as a document.
DOC_QUAD_MIN_AREA = 0.5
# Regex: a line that begins with a number marker like "1." or "12)"
LIST_MARKER_RE = re.compile(r"^\s*(\d{1,2})[.)]\s+")
# Regex: a line that is the wrapped continuation of the previous item
# (starts with punctuation, not a real word).
CONTINUATION_START_RE = re.compile(r"^[.,:;፡፥፣]")
# Regex: a line that ends like a real list item (Amharic full stop or '?').
LIST_ITEM_END_RE = re.compile(r"(?:[?]|።|፡፡|\?)$")

app = FastAPI(title="OCR Pipeline API", version="5.0.0")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


#  Geometry helpers


def encode_image(img: np.ndarray) -> str:
    _, buf = cv2.imencode(".jpg", img, [cv2.IMWRITE_JPEG_QUALITY, 80])
    return base64.b64encode(buf.tobytes()).decode("utf-8")


def order_points(pts):
    rect = np.zeros((4, 2), dtype="float32")
    s = pts.sum(axis=1)
    rect[0] = pts[np.argmin(s)]
    rect[2] = pts[np.argmax(s)]
    diff = np.diff(pts, axis=1)
    rect[1] = pts[np.argmin(diff)]
    rect[3] = pts[np.argmax(diff)]
    return rect


def four_point_transform(image, pts):
    rect = order_points(pts)
    tl, tr, br, bl = rect
    maxW = max(int(np.linalg.norm(br - bl)), int(np.linalg.norm(tr - tl)))
    maxH = max(int(np.linalg.norm(tr - br)), int(np.linalg.norm(tl - bl)))
    dst = np.array(
        [[0, 0], [maxW - 1, 0], [maxW - 1, maxH - 1], [0, maxH - 1]],
        dtype="float32",
    )
    M = cv2.getPerspectiveTransform(rect, dst)
    return cv2.warpPerspective(image, M, (maxW, maxH))


def find_document_contour(edged, orig_area):
    """
    Return the 4 corners of the document page only if a clear quadrilateral
    covering most of the image is found. Bad guesses make perspective warp
    destroy the text, so we err on the side of "no detection".
    """
    contours, _ = cv2.findContours(edged, cv2.RETR_LIST, cv2.CHAIN_APPROX_SIMPLE)
    for c in sorted(contours, key=cv2.contourArea, reverse=True)[:10]:
        area = cv2.contourArea(c)
        if area < orig_area * DOC_QUAD_MIN_AREA:
            continue
        peri = cv2.arcLength(c, True)
        approx = cv2.approxPolyDP(c, 0.02 * peri, True)
        if len(approx) == 4:
            return approx.reshape(4, 2).astype("float32")
    return None


def deskew(gray):
    thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY_INV | cv2.THRESH_OTSU)[1]
    coords = np.column_stack(np.where(thresh > 0))
    if len(coords) < 10:
        return gray
    angle = cv2.minAreaRect(coords)[-1]
    angle = -(90 + angle) if angle < -45 else -angle
    if abs(angle) < 0.3:
        return gray
    h, w = gray.shape
    M = cv2.getRotationMatrix2D((w // 2, h // 2), angle, 1.0)
    return cv2.warpAffine(
        gray, M, (w, h), flags=cv2.INTER_CUBIC, borderMode=cv2.BORDER_REPLICATE
    )


#  Preprocessing for printed Amharic


def preprocess_for_ocr(bgr: np.ndarray) -> tuple[np.ndarray, np.ndarray]:
    """
    Printed-text pipeline tuned for Amharic — feeds Tesseract LSTM grayscale.

      grayscale → upscale → deskew → CLAHE → mild bilateral denoise.

    Returns (enhanced_gray_for_ocr, preview_binary). Modern Tesseract 4/5 LSTM
    (OEM 1) recognises grayscale far better than a pre-binarised image, because
    binarisation often destroys the thin Ethiopic strokes. The binary image is
    only kept for display in the pipeline visualisation.
    """
    gray = cv2.cvtColor(bgr, cv2.COLOR_BGR2GRAY)
    h, w = gray.shape

    long_edge = max(h, w)
    if long_edge < OCR_TARGET_PX:
        scale = OCR_TARGET_PX / long_edge
        gray = cv2.resize(
            gray, (int(w * scale), int(h * scale)), interpolation=cv2.INTER_CUBIC
        )
    elif long_edge > OCR_TARGET_PX * 1.5:
        scale = OCR_TARGET_PX / long_edge
        gray = cv2.resize(
            gray, (int(w * scale), int(h * scale)), interpolation=cv2.INTER_AREA
        )

    gray = deskew(gray)

    clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
    gray = clahe.apply(gray)

    # Light denoise that preserves stroke edges; Tesseract is happy with this.
    enhanced = cv2.bilateralFilter(gray, 5, 50, 50)

    # Binary preview only — not used for OCR.
    _, preview_bin = cv2.threshold(
        enhanced, 0, 255, cv2.THRESH_BINARY | cv2.THRESH_OTSU
    )

    return enhanced, preview_bin


#  Text post-processing


def _token_is_valid(token: str) -> bool:
    if not token:
        return False
    has_ethiopic = bool(ETHIOPIC_RE.search(token))
    has_latin = len(re.findall(r"[a-zA-Z]", token)) >= 2
    has_digit = any(c.isdigit() for c in token)
    if not (has_ethiopic or has_latin or has_digit):
        return False
    real = len(REAL_CHAR_RE.findall(token))
    total = sum(1 for c in token if not c.isspace())
    return total > 0 and real / total >= 0.4


def _is_likely_garbage_line(stripped: str) -> bool:
    """
    Detect short, mixed-script lines that almost always come from edge noise
    or watermarks (e.g. "aA ኩን", "oO mH ው"). Real Amharic lines are usually
    > 10 chars and either pure Ethiopic or have long-enough words.
    """
    if len(stripped) > 12:
        return False
    has_latin = bool(re.search(r"[a-zA-Z]", stripped))
    has_ethiopic = bool(ETHIOPIC_RE.search(stripped))
    if has_latin and has_ethiopic:
        return True
    tokens = stripped.split()
    if tokens and all(len(t) <= 2 for t in tokens):
        return True
    return False


def clean_ocr_text(text: str) -> str:
    """Drop garbage lines but keep legitimate ones even if a few tokens are bad."""
    out_lines = []
    for line in text.split("\n"):
        line = DOT_RUN_RE.sub(" ", line)
        line = SPACE_RUN_RE.sub(" ", line).strip()
        if not line:
            continue
        if _is_likely_garbage_line(line):
            continue
        tokens = line.split()
        good_tokens = [t for t in tokens if _token_is_valid(t)]
        if good_tokens and len(good_tokens) >= max(1, len(tokens) // 3):
            out_lines.append(" ".join(good_tokens))
    return "\n".join(out_lines)


def score_ocr_text(cleaned: str) -> int:
    score = 0
    for token in cleaned.split():
        if ETHIOPIC_RE.search(token) and len(token) >= 2:
            score += 1
        elif LATIN_WORD_RE.search(token):
            score += 1
    return score


def resolve_ocr_lang() -> str:
    try:
        langs = pytesseract.get_languages(config="")
    except Exception:
        return "eng"
    if "amh" in langs:
        return OCR_LANG if "eng" in langs else "amh"
    logger.warning("Amharic (amh) not installed; falling back to eng")
    return "eng"


#  OCR — single fast pass


def _ocr_pass(pil_img: PILImage.Image, lang: str, psm: int) -> tuple[str, int, int]:
    """Single Tesseract pass; returns (cleaned_text, word_count, avg_conf).

    Drops lines whose average word confidence is below MIN_LINE_CONF so that
    edge noise at the bottom of a page (e.g. "aA & W ኩን") is suppressed
    without losing legitimate low-confidence single tokens elsewhere.
    """
    cfg = f"--oem 1 --psm {psm} -c preserve_interword_spaces=1"
    data = pytesseract.image_to_data(
        pil_img, config=cfg, lang=lang, output_type=pytesseract.Output.DICT
    )

    line_words: dict[tuple, list[str]] = {}
    line_confs: dict[tuple, list[int]] = {}
    for i, word in enumerate(data["text"]):
        if not word.strip():
            continue
        conf = int(data["conf"][i])
        if conf < MIN_WORD_CONF:
            continue
        key = (
            data["block_num"][i],
            data["par_num"][i],
            data["line_num"][i],
        )
        line_words.setdefault(key, []).append(word)
        line_confs.setdefault(key, []).append(conf)

    # Drop low-confidence lines (typical of edge/footer noise)
    good_keys = [
        k for k in line_words
        if sum(line_confs[k]) / max(len(line_confs[k]), 1) >= MIN_LINE_CONF
    ]

    raw_lines = [" ".join(line_words[k]) for k in sorted(good_keys)]
    raw = "\n".join(raw_lines)
    cleaned = clean_ocr_text(raw)
    cleaned = _infer_list_numbering(cleaned)

    all_confs = [c for k in good_keys for c in line_confs[k]]
    avg_conf = int(np.mean(all_confs)) if all_confs else 0
    return cleaned, score_ocr_text(cleaned), avg_conf


def _infer_list_numbering(text: str) -> str:
    """
    Tesseract's amh+eng LSTM often recognises the first numbered list marker
    (``1.``) but drops the remaining digits 2..N. When we see a sequential
    list that breaks after the first item, we re-number the consecutive
    question/item lines that follow.

    Conservative: only triggers when the very next non-empty line lacks a
    marker AND looks like a list item (ends with ``?`` / ``፡፡`` / ``።``).
    """
    lines = text.split("\n")
    out: list[str] = []
    i = 0
    while i < len(lines):
        line = lines[i]
        m = LIST_MARKER_RE.match(line)
        if not m:
            out.append(line)
            i += 1
            continue

        start_n = int(m.group(1))
        out.append(line)
        counter = start_n + 1
        j = i + 1
        while j < len(lines):
            nxt = lines[j]
            nxt_stripped = nxt.strip()
            if not nxt_stripped:
                break  # blank line ends the list paragraph
            if LIST_MARKER_RE.match(nxt_stripped):
                break  # Tesseract already recognised this marker
            is_continuation = (
                CONTINUATION_START_RE.match(nxt_stripped)
                # Short fragment that doesn't end like a real list item is
                # almost certainly a wrapped continuation.
                or (
                    len(nxt_stripped) < 20
                    and not LIST_ITEM_END_RE.search(nxt_stripped)
                )
            )
            if is_continuation:
                out[-1] = out[-1] + " " + nxt_stripped
                j += 1
                continue
            out.append(f"{counter}. {nxt_stripped}")
            counter += 1
            j += 1
        i = j

    return "\n".join(out)


def run_ocr(pil_img: PILImage.Image, lang: str) -> tuple[str, int, int]:
    """
    Run Tesseract LSTM, picking the best layout mode:
      - PSM 3: fully automatic (best for multi-column / mixed layout)
      - PSM 4: single column of variable text sizes (fallback when PSM 3
               merges or misses sections)

    Skips PSM 4 when PSM 3 already produced a strong, high-confidence result
    to keep typical pages fast (~3-4 s instead of 8-10 s).
    """
    text, score, conf = _ocr_pass(pil_img, lang, psm=3)
    logger.info(f"PSM 3 OEM 1 ({lang}): {score} words, conf={conf}%")

    if score >= 40 and conf >= 75:
        logger.info("PSM 3 confident — skipping PSM 4 fallback")
        return text, score, conf

    text2, score2, conf2 = _ocr_pass(pil_img, lang, psm=4)
    logger.info(f"PSM 4 OEM 1 ({lang}): {score2} words, conf={conf2}%")

    if score2 * max(conf2, 1) > score * max(conf, 1):
        return text2, score2, conf2
    return text, score, conf


#  Main pipeline


def _resize_for_preview(img: np.ndarray) -> np.ndarray:
    h, w = img.shape[:2]
    if max(h, w) <= PREVIEW_PX:
        return img
    s = PREVIEW_PX / max(h, w)
    return cv2.resize(img, (int(w * s), int(h * s)), interpolation=cv2.INTER_AREA)


def run_ocr_pipeline(img_bytes: bytes, include_pipeline: bool = True) -> dict:
    t0 = time.perf_counter()

    nparr = np.frombuffer(img_bytes, np.uint8)
    original = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
    if original is None:
        raise ValueError("Could not decode image")

    h, w = original.shape[:2]
    display = _resize_for_preview(original)
    dh, dw = display.shape[:2]
    orig_area = dh * dw

    pipeline = {}
    if include_pipeline:
        pipeline["original"] = encode_image(display)

    gray_d = cv2.cvtColor(display, cv2.COLOR_BGR2GRAY)
    blurred = cv2.GaussianBlur(gray_d, (5, 5), 0)
    edges = cv2.Canny(blurred, 20, 80)
    dilated = cv2.dilate(edges, np.ones((5, 5), np.uint8), iterations=2)
    if include_pipeline:
        pipeline["canny"] = encode_image(cv2.cvtColor(dilated, cv2.COLOR_GRAY2BGR))

    doc_corners = find_document_contour(dilated, orig_area)
    if include_pipeline:
        contour_img = display.copy()
        if doc_corners is not None:
            cv2.drawContours(contour_img, [doc_corners.astype(int)], -1, (0, 255, 80), 3)
            for pt in doc_corners:
                cv2.circle(contour_img, tuple(pt.astype(int)), 12, (0, 100, 255), -1)
        else:
            cnts, _ = cv2.findContours(dilated, cv2.RETR_LIST, cv2.CHAIN_APPROX_SIMPLE)
            cv2.drawContours(
                contour_img,
                sorted(cnts, key=cv2.contourArea, reverse=True)[:5],
                -1,
                (0, 255, 80),
                2,
            )
        pipeline["contours"] = encode_image(contour_img)

    if doc_corners is not None:
        warped = four_point_transform(original, doc_corners * (max(h, w) / max(dh, dw)))
    else:
        warped = original

    if include_pipeline:
        pipeline["warped"] = encode_image(_resize_for_preview(warped))

    t_pre = time.perf_counter()
    enhanced_gray, preview_bin = preprocess_for_ocr(warped)
    if include_pipeline:
        pipeline["threshold"] = encode_image(
            cv2.cvtColor(_resize_for_preview(preview_bin), cv2.COLOR_GRAY2BGR)
        )

    lang = resolve_ocr_lang()
    pil_img = PILImage.fromarray(enhanced_gray)

    t_ocr = time.perf_counter()
    text, word_count, avg_conf = run_ocr(pil_img, lang)
    t_end = time.perf_counter()

    logger.info(
        f"Timings — preprocess: {t_pre - t0:.2f}s, "
        f"prep_to_ocr: {t_ocr - t_pre:.2f}s, "
        f"tesseract: {t_end - t_ocr:.2f}s, "
        f"total: {t_end - t0:.2f}s"
    )

    return {
        "pipeline_images": pipeline if include_pipeline else None,
        "extracted_text": text,
        "word_count": word_count,
        "avg_confidence": avg_conf,
        "doc_detected": doc_corners is not None,
        "image_size": {"width": w, "height": h},
        "ocr_lang": lang,
        "elapsed_ms": int((t_end - t0) * 1000),
    }


#  DOCX export


def combine_pages_text(pages: list[dict]) -> str:
    parts = []
    for p in pages:
        n = p.get("page", len(parts) + 1)
        text = (p.get("extracted_text") or "").strip()
        parts.append(f"--- ገጽ {n} / Page {n} ---\n{text}")
    return "\n\n".join(parts)


def build_docx_bytes(title: str, pages: list[dict]) -> bytes:
    doc = Document()
    doc.core_properties.title = title

    heading = doc.add_heading(title, level=0)
    for run in heading.runs:
        run.font.name = DOCX_FONT
        run.font.size = Pt(18)

    for i, page in enumerate(pages):
        if i > 0:
            doc.add_page_break()
        page_num = page.get("page", i + 1)
        label = page.get("title") or f"ገጽ {page_num} / Page {page_num}"
        h = doc.add_heading(label, level=1)
        for run in h.runs:
            run.font.name = DOCX_FONT

        text = (page.get("text") or page.get("extracted_text") or "").strip()
        para = doc.add_paragraph(text)
        for run in para.runs:
            run.font.name = DOCX_FONT
            run.font.size = Pt(12)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


#  API models


class DocxPageIn(BaseModel):
    page: int = 1
    title: Optional[str] = None
    text: str = ""


class DocxExportRequest(BaseModel):
    title: str = "Amharic Journal OCR"
    pages: List[DocxPageIn] = Field(..., min_length=1)


#  Endpoints


@app.get("/")
async def root():
    return {
        "service": "OCR Pipeline API",
        "version": "5.0.0",
        "ocr_lang": OCR_LANG,
        "features": ["amharic", "batch", "docx"],
    }


@app.get("/health")
async def health():
    try:
        langs = pytesseract.get_languages(config="")
        return {
            "status": "ok",
            "tesseract": str(pytesseract.get_tesseract_version()),
            "languages": langs,
            "amharic_ready": "amh" in langs,
            "ocr_lang": resolve_ocr_lang(),
        }
    except Exception as e:
        return {"status": "degraded", "error": str(e)}


@app.post("/ocr")
async def ocr_endpoint(
    file: UploadFile = File(...),
    include_pipeline: bool = Query(True),
):
    if not file.content_type or not file.content_type.startswith("image/"):
        raise HTTPException(400, "File must be an image")
    data = await file.read()
    if len(data) > 25 * 1024 * 1024:
        raise HTTPException(413, "Image too large (max 25 MB)")
    try:
        return run_ocr_pipeline(data, include_pipeline=include_pipeline)
    except ValueError as e:
        raise HTTPException(400, str(e))
    except Exception as e:
        logger.error(f"Pipeline error: {e}", exc_info=True)
        raise HTTPException(500, f"Pipeline error: {str(e)}")


@app.post("/ocr/batch")
async def ocr_batch_endpoint(files: List[UploadFile] = File(...)):
    if not files:
        raise HTTPException(400, "At least one image is required")
    if len(files) > MAX_BATCH_PAGES:
        raise HTTPException(400, f"Maximum {MAX_BATCH_PAGES} images per batch")

    pages: list[dict] = []
    errors: list[dict] = []
    pipeline_images = None

    for idx, file in enumerate(files):
        if not file.content_type or not file.content_type.startswith("image/"):
            errors.append({"page": idx + 1, "error": "Not an image"})
            continue
        data = await file.read()
        if len(data) > 25 * 1024 * 1024:
            errors.append({"page": idx + 1, "error": "Image too large (max 25 MB)"})
            continue
        try:
            include_pipeline = pipeline_images is None
            result = run_ocr_pipeline(data, include_pipeline=include_pipeline)
            if include_pipeline:
                pipeline_images = result.get("pipeline_images")
            pages.append(
                {
                    "page": idx + 1,
                    "filename": file.filename,
                    "extracted_text": result["extracted_text"],
                    "word_count": result["word_count"],
                    "avg_confidence": result["avg_confidence"],
                    "doc_detected": result["doc_detected"],
                    "ocr_lang": result["ocr_lang"],
                }
            )
        except Exception as e:
            logger.error(f"Batch page {idx + 1} error: {e}", exc_info=True)
            errors.append({"page": idx + 1, "error": str(e)})

    if not pages:
        raise HTTPException(400, "No pages processed successfully")

    combined_text = combine_pages_text(pages)
    total_words = sum(p["word_count"] for p in pages)
    confs = [p["avg_confidence"] for p in pages if p["avg_confidence"] > 0]
    avg_conf = int(np.mean(confs)) if confs else 0

    return {
        "mode": "batch",
        "page_count": len(pages),
        "pages": pages,
        "combined_text": combined_text,
        "total_word_count": total_words,
        "avg_confidence": avg_conf,
        "ocr_lang": pages[0]["ocr_lang"],
        "pipeline_images": pipeline_images,
        "doc_detected": any(p["doc_detected"] for p in pages),
        "errors": errors or None,
    }


@app.post("/ocr/docx")
async def export_docx_endpoint(body: DocxExportRequest, as_file: bool = Query(False)):
    pages = [
        {"page": p.page, "title": p.title, "text": p.text} for p in body.pages
    ]
    try:
        docx_bytes = build_docx_bytes(body.title, pages)
    except Exception as e:
        logger.error(f"DOCX build error: {e}", exc_info=True)
        raise HTTPException(500, f"Could not create Word document: {str(e)}")

    filename = "journal_ocr.docx"
    if as_file:
        return Response(
            content=docx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )

    return {
        "filename": filename,
        "docx_base64": base64.b64encode(docx_bytes).decode("utf-8"),
        "size_bytes": len(docx_bytes),
    }
